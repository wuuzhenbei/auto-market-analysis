"""
生成Word格式简历
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_resume():
    doc = Document()

    # 设置字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # ====== 标题 ======
    title = doc.add_heading('王子柏', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('数据分析实习')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 102, 204)

    # ====== 基本信息 ======
    doc.add_heading('基本信息', level=1)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    info = [
        ('手机', '182-4395-8996', '邮箱', 'wangzibo720@gmail.com'),
        ('学校', '大连民族大学', '专业', '物联网工程（本科）'),
        ('毕业', '2027年6月', '求职意向', '数据分析实习'),
        ('城市', '上海', '到岗时间', '随时，6个月+，每周5天'),
    ]

    for i, (k1, v1, k2, v2) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = k1
        row.cells[1].text = v1
        row.cells[2].text = k2
        row.cells[3].text = v2

    # ====== 个人优势 ======
    doc.add_heading('个人优势', level=1)

    advantages = [
        '独立完成从零到一的数据平台项目，具备数据采集、处理、分析、可视化的全链路实操能力',
        '对数据驱动业务决策有浓厚兴趣，善于用AI工具提升数据处理效率',
        '物联网工程背景训练了对数据从产生到处理全链路的系统理解',
        '学习能力强，Python、Vue、AI Agent均为项目驱动自学，能快速上手新技术栈',
        '具备工程化思维，项目包含定时任务、数据校验、自动化部署等生产级功能',
    ]

    for adv in advantages:
        p = doc.add_paragraph(adv, style='List Bullet')

    # ====== 项目经历 ======
    doc.add_heading('项目经历', level=1)

    # 项目1
    doc.add_heading('汽车市场数据分析平台 | 独立开发 | 2026.08', level=2)

    p = doc.add_paragraph()
    p.add_run('项目描述：').bold = True
    p.add_run('独立完成汽车行业数据分析平台，实现从数据采集到可视化展示的完整数据链路，覆盖1000+车型、26个品牌、7个月度销量数据。')

    p = doc.add_paragraph()
    p.add_run('技术栈：').bold = True
    p.add_run('Python、SQLite、Pandas、Plotly、Streamlit、Cloudflare Tunnel、Git')

    p = doc.add_paragraph()
    p.add_run('工作内容：').bold = True

    work_items = [
        '数据采集模块：使用browser-harness实现浏览器自动化采集，研究并破解懂车帝反爬虫机制，通过Cookie绕过验证，实现70+条销量数据稳定采集',
        '数据存储与处理：设计SQLite数据库schema（6张表），使用Pandas进行数据清洗、转换、聚合，实现品牌自动分类（自主/合资/豪华/新势力）',
        '数据分析与可视化：实现7个分析维度（市场总览、品牌、价格、新能源、口碑、趋势、城市），使用Plotly生成20+交互式图表',
        '工程化与自动化：设计定时任务系统（月度采集、日度校验、月度导出），实现数据校验模块，使用Cloudflare Tunnel实现公网访问',
    ]

    for item in work_items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('项目成果：').bold = True

    results = [
        '独立完成从0到1的数据平台搭建',
        '实现懂车帝反爬虫破解，稳定采集7个月度数据',
        '构建完整的数据分析Dashboard，支持7个分析维度',
        '项目已开源至GitHub：https://github.com/wuuzhenbei/auto-market-analysis',
    ]

    for result in results:
        doc.add_paragraph(result, style='List Bullet')

    # 项目2
    doc.add_heading('灵讯IM即时通讯系统 | 参与开发 | 2026.07', level=2)

    p = doc.add_paragraph()
    p.add_run('项目描述：').bold = True
    p.add_run('企业级即时通讯系统，支持文字、图片、文件传输，实现Docker容器化部署。')

    p = doc.add_paragraph()
    p.add_run('技术栈：').bold = True
    p.add_run('Docker、Cloudflare Tunnel、Nginx、WebSocket')

    p = doc.add_paragraph()
    p.add_run('工作内容：').bold = True

    work_items2 = [
        '参与系统架构设计与部署',
        '使用Docker实现容器化部署',
        '配置Cloudflare Tunnel实现公网访问',
        '实现域名解析与HTTPS配置',
    ]

    for item in work_items2:
        doc.add_paragraph(item, style='List Bullet')

    # ====== 技能清单 ======
    doc.add_heading('技能清单', level=1)

    skills = [
        ('编程语言', 'Python（熟练）、JavaScript/Vue（熟练）、SQL（熟练）'),
        ('数据分析', 'Pandas（熟练）、NumPy（熟练）、Matplotlib/Seaborn/Plotly（熟练）、SQLite/MySQL（熟练）'),
        ('工具平台', 'Git（熟练）、Docker（熟练）、Cloudflare（熟练）、Streamlit（熟练）、browser-harness（熟练）'),
        ('AI工具', 'Claude/GPT（熟练）、Cursor/Claude Code（熟练）'),
        ('其他技能', '网络爬虫与反爬虫技术、数据采集与处理、自动化脚本编写、API接口调用'),
    ]

    for skill, detail in skills:
        p = doc.add_paragraph()
        p.add_run(f'{skill}：').bold = True
        p.add_run(detail)

    # ====== 教育背景 ======
    doc.add_heading('教育背景', level=1)

    p = doc.add_paragraph()
    p.add_run('大连民族大学 | 物联网工程（本科） | 2022.09 - 2027.06').bold = True

    p = doc.add_paragraph()
    p.add_run('主修课程：').bold = True
    p.add_run('数据结构、数据库原理、Python程序设计、传感器技术、数据采集与处理、计算机网络、操作系统')

    # ====== 自我评价 ======
    doc.add_heading('自我评价', level=1)

    evaluations = [
        '数据驱动：对数据分析有浓厚兴趣，善于从数据中发现业务洞察',
        '工程化思维：注重代码质量、系统设计、自动化流程',
        '快速学习：Python、Vue、AI Agent均为项目驱动自学，能快速上手新技术栈',
        '独立解决问题：具备独立完成从0到1项目的能力',
        'AI赋能：善于利用AI工具提升开发效率，具备AI Agent开发经验',
    ]

    for eval in evaluations:
        doc.add_paragraph(eval, style='List Bullet')

    # ====== 附件 ======
    doc.add_heading('附件', level=1)

    p = doc.add_paragraph()
    p.add_run('GitHub项目地址：').bold = True
    p.add_run('https://github.com/wuuzhenbei/auto-market-analysis')

    p = doc.add_paragraph()
    p.add_run('项目在线演示：').bold = True
    p.add_run('https://autocar.050311.xyz')

    # 保存
    doc.save('E:/auto-market-analysis/resume/王子柏_数据分析实习简历.docx')
    print("简历已生成: E:/auto-market-analysis/resume/王子柏_数据分析实习简历.docx")


if __name__ == "__main__":
    create_resume()
